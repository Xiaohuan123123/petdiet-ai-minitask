"""
将 PRD Markdown 文件转换为 Word 文档
"""
import re
import sys
import io
from pathlib import Path

# Windows 编码修复
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_markdown_to_doc(doc, md_text):
    """将 Markdown 文本解析并添加到 Word 文档"""
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 水平线
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = clean_markdown(text)
            heading = doc.add_heading(text, level=min(level, 3))
            set_heading_font(heading, level)
            i += 1
            continue

        # 表格
        if '|' in line and i + 1 < len(lines) and '|---' in lines[i + 1]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1

            # 解析表格
            add_table_to_doc(doc, table_lines)
            continue

        # 有序列表
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if ol_match:
            list_items = []
            while i < len(lines):
                m = re.match(r'^(\d+)\.\s+(.+)$', lines[i].strip())
                if not m:
                    break
                list_items.append(clean_markdown(m.group(2)))
                i += 1

            for idx, item in enumerate(list_items):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(f"{idx + 1}. ")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
                run.font.bold = True
                add_formatted_text(p, item)
            continue

        # 无序列表
        ul_match = re.match(r'^[-*]\s+(.+)$', line.strip()) or re.match(r'^>\s+(.+)$', line.strip())
        if ul_match:
            list_items = []
            while i < len(lines):
                m = re.match(r'^[-*]\s+(.+)$', lines[i].strip()) or re.match(r'^>\s+(.+)$', lines[i].strip())
                if not m:
                    break
                list_items.append(clean_markdown(m.group(1)))
                i += 1

            for item in list_items:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run("• ")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
                add_formatted_text(p, item)
            continue

        # 引用块（以 > 开头但不是无序列表的分隔符）
        if line.strip().startswith('> '):
            quote_text = clean_markdown(line.strip()[2:])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(quote_text)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            run.font.italic = True
            i += 1
            continue

        # 粗体标题（如 **核心数据速览**：）
        bold_title = re.match(r'^\*\*(.+?)\*\*[：:]?\s*(.*)$', line)
        if bold_title:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(bold_title.group(1))
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x45, 0x1A, 0x03)
            if bold_title.group(2):
                add_formatted_text(p, bold_title.group(2))
            i += 1
            continue

        # 普通段落
        text = clean_markdown(line)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_formatted_text(p, text, size=11)
        i += 1


def add_table_to_doc(doc, lines):
    """添加表格到文档"""
    if len(lines) < 2:
        return

    # 解析表头
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]

    # 跳过分隔行 (|---|---|)
    data_start = 1
    if re.match(r'^[\|\s\-:]+$', lines[1]):
        data_start = 2

    # 解析数据行
    rows = []
    for line in lines[data_start:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)

    if not headers:
        return

    # 创建表格
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(clean_markdown(header))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        set_cell_shading(cell, 'FEF3C7')

    # 填充数据行
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < len(headers):
                cell = table.rows[i + 1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_formatted_text(p, clean_markdown(cell_text), size=10)

    # 添加空行分隔
    doc.add_paragraph()


def clean_markdown(text):
    """清理 Markdown 格式标记（保留加粗等）"""
    return text


def add_formatted_text(paragraph, text, size=11, color=None):
    """添加带格式的文本到段落，处理 **加粗** 标记"""
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0x45, 0x1A, 0x03)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            if color:
                run.font.color.rgb = color


def set_heading_font(heading, level):
    """设置标题字体"""
    run = heading.runs[0] if heading.runs else None
    if not run:
        return

    if level == 0:  # Title
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    elif level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x45, 0x1A, 0x03)


def main():
    prd_file = Path(__file__).parent.parent / 'competitive-analysis-v1.md'
    output_file = Path(__file__).parent.parent / '宠食记-PetDiet-AI-竞品分析-v1.docx'

    if not prd_file.exists():
        print(f"[错误] 未找到文件: {prd_file}")
        sys.exit(1)

    print(f"[读取] {prd_file}")
    with open(prd_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 跳过 YAML front matter
    md_text = re.sub(r'^---.*?---\s*', '', md_text, flags=re.DOTALL)

    print("[生成] Word文档...")
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_markdown_to_doc(doc, md_text)

    # 设置所有表格的宽度
    for table in doc.tables:
        table.autofit = True

    doc.save(str(output_file))
    print(f"[完成] 已生成: {output_file}")


if __name__ == '__main__':
    main()
